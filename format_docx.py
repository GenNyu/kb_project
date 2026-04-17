#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
import re
from pathlib import Path
from typing import Any, Iterable, Optional

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - runtime dependency check
    print("Thiếu thư viện python-docx. Cài bằng: pip install -r requirements.txt", file=sys.stderr)
    raise SystemExit(2) from exc


def add_label(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True


def add_text(doc: Document, text: Optional[str]) -> None:
    if text is None:
        return
    clean = text.strip()
    if not clean:
        return
    doc.add_paragraph(clean)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for raw in items:
        item = str(raw).strip()
        if not item:
            continue
        try:
            doc.add_paragraph(item, style="List Bullet")
        except KeyError:
            doc.add_paragraph(f"- {item}")


def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Chuẩn hoá nội dung thành DOCX theo mẫu A-E cho Control Objective."
    )
    parser.add_argument(
        "input_json",
        type=Path,
        help="JSON hoặc DOCX nguồn. Có thể là thư mục để xử lý hàng loạt.",
    )
    parser.add_argument(
        "output_docx",
        type=Path,
        help="File DOCX xuất ra hoặc thư mục đích khi xử lý hàng loạt.",
    )
    return parser.parse_args(argv)


def load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        print(f"Không tìm thấy file: {path}", file=sys.stderr)
        raise SystemExit(2)
    if path.suffix.lower() == ".docx":
        return build_data_from_docx(path)
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except UnicodeDecodeError:
        print("File đầu vào không phải UTF-8 JSON. Nếu là DOCX, hãy dùng đuôi .docx.", file=sys.stderr)
        raise SystemExit(2)
    except json.JSONDecodeError as exc:
        print(f"Lỗi JSON: {exc}", file=sys.stderr)
        raise SystemExit(2)


def extract_ids_from_docx(doc: Document) -> tuple[str, str]:
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    full_text = "\n".join(paragraphs)
    match = re.search(r"\b(\d+)\.(\d+)\b", full_text)
    requirement_id = match.group(1) if match else ""
    control_objective_id = match.group(0) if match else ""
    return requirement_id, control_objective_id


def clone_paragraph(dst_doc: Document, src_paragraph) -> None:
    para = dst_doc.add_paragraph()
    try:
        para.style = src_paragraph.style
    except Exception:
        pass
    try:
        para.alignment = src_paragraph.alignment
    except Exception:
        pass
    try:
        para.paragraph_format.left_indent = src_paragraph.paragraph_format.left_indent
        para.paragraph_format.right_indent = src_paragraph.paragraph_format.right_indent
        para.paragraph_format.first_line_indent = src_paragraph.paragraph_format.first_line_indent
        para.paragraph_format.space_before = src_paragraph.paragraph_format.space_before
        para.paragraph_format.space_after = src_paragraph.paragraph_format.space_after
        para.paragraph_format.line_spacing = src_paragraph.paragraph_format.line_spacing
    except Exception:
        pass

    if not src_paragraph.runs:
        para.add_run(src_paragraph.text)
        return

    for run in src_paragraph.runs:
        new_run = para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        try:
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
        except Exception:
            pass


def build_docx_from_docx(input_docx: Path, output_docx: Path) -> None:
    src_doc = Document(input_docx)
    requirement_id, control_objective_id = extract_ids_from_docx(src_doc)

    doc = Document()

    add_label(doc, f"A. Tài liệu gốc của Requirement {requirement_id}".strip())

    add_label(doc, f"B. Summary Overview của Control Objective {control_objective_id}".strip())
    base = (
        f"Tài liệu này mô tả chi tiết Control Objective {control_objective_id} "
        f"của Requirement {requirement_id} trong PCI-DSS v4.0.1,"
    )
    base = base + " tập trung vào việc"
    base = base + "\n" + "Mục tiêu là"
    base = base + "\n" + "Gồm sub-requirement chính:\n"
    add_text(doc, base)

    add_label(doc, f"C. Key Points của Control Objective {control_objective_id}".strip())

    add_label(doc, f"D. Deep Summary của Control Objective {control_objective_id}".strip())
    add_text(doc, "Bối cảnh:")
    add_text(doc, "Nội dung cốt lõi:")
    add_text(doc, "Dữ liệu đáng chú ý:")
    add_text(doc, "Rủi ro / Lưu ý:")

    add_label(doc, f"E. Structured Output của Control Objective {control_objective_id}".strip())
    for p in src_doc.paragraphs:
        clone_paragraph(doc, p)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def build_docx(data: dict[str, Any], output_docx: Path) -> None:
    requirement_id = str(data.get("requirement_id", "")).strip()
    control_objective_id = str(data.get("control_objective_id", "")).strip()
    sub_req_count = data.get("sub_requirement_count")
    summary_overview = data.get("summary_overview")
    overview_sentence = data.get("overview_sentence")
    focus = data.get("focus")
    key_points = data.get("key_points", [])
    deep_summary = data.get("deep_summary", {})
    structured_output = data.get("structured_output")
    cut_sub_requirements = data.get("cut_sub_requirements", [])

    doc = Document()

    # A
    add_label(doc, f"A. Tài liệu gốc của Requirement {requirement_id}".strip())

    # B
    add_label(doc, f"B. Summary Overview của Control Objective {control_objective_id}".strip())
    if overview_sentence:
        add_text(doc, overview_sentence)
    else:
        base = (
            f"Tài liệu này mô tả chi tiết Control Objective {control_objective_id} "
            f"của Requirement {requirement_id} trong PCI-DSS v4.0.1."
        )
        base = base.rstrip(",") + ", tập trung vào việc"
        base = base + "\n" + "Mục tiêu là"
        base = base + "\n" + f"Gồm  sub-requirement chính:"
        add_text(doc, base)
    add_text(doc, summary_overview)


    # C
    add_label(doc, f"C. Key Points của Control Objective {control_objective_id}".strip())
    add_bullets(doc, key_points)

    # D
    add_label(doc, f"D. Deep Summary của Control Objective {control_objective_id}".strip())
    add_text(doc, "Bối cảnh:")
    add_text(doc, deep_summary.get("context"))
    add_text(doc, "Nội dung cốt lõi:")
    add_text(doc, deep_summary.get("core"))
    add_text(doc, "Dữ liệu đáng chú ý:")
    add_text(doc, deep_summary.get("notable_data"))
    add_text(doc, "Rủi ro / Lưu ý:")
    add_text(doc, deep_summary.get("risks"))

    # E
    add_label(doc, f"E. Structured Output của Control Objective {control_objective_id}".strip())
    add_text(doc, structured_output)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main(argv: Optional[list[str]] = None) -> int:
    args = parse_args(argv)
    if args.input_json.is_dir():
        if args.output_docx.exists() and not args.output_docx.is_dir():
            print("Output phải là thư mục khi input là thư mục", file=sys.stderr)
            return 2
        inputs = sorted(
            p
            for p in args.input_json.iterdir()
            if p.suffix.lower() in {".json", ".docx"}
        )
        if not inputs:
            print("Không tìm thấy file .json hoặc .docx trong thư mục", file=sys.stderr)
            return 2
        args.output_docx.mkdir(parents=True, exist_ok=True)
        for input_path in inputs:
            out_name = f"{input_path.stem}.docx"
            if input_path.suffix.lower() == ".docx":
                build_docx_from_docx(input_path, args.output_docx / out_name)
            else:
                data = load_json(input_path)
                build_docx(data, args.output_docx / out_name)
    else:
        if args.output_docx.is_dir():
            args.output_docx.mkdir(parents=True, exist_ok=True)
            out_path = args.output_docx / f"{args.input_json.stem}.docx"
        else:
            out_path = args.output_docx
        if args.input_json.suffix.lower() == ".docx":
            build_docx_from_docx(args.input_json, out_path)
        else:
            data = load_json(args.input_json)
            build_docx(data, out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
