import os
import re
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement

# =========================
# CONFIG
# =========================
INPUT_FOLDER = "format"
OUTPUT_FOLDER = "reformat"

TARGET_FIELDS = [
    "Defined Approach Requirements",
    "Customized Approach Objective",
    "Applicability Notes",
    "Defined Approach Testing Procedures",
    "Purpose",
    "Good Practice",
    "Guidance - Purpose",
    "Guidance - Good Practice",
    "Guidance - Definitions",
    "Guidance - Examples",
    "Guidance - Further Information",
    "Definitions",
    "Examples",
    "Further Information",
]

# sort dài trước để tránh match nhầm
TARGET_FIELDS = sorted(TARGET_FIELDS, key=len, reverse=True)

SUBREQ_PATTERN = re.compile(r"^\s*Sub-requirement:\s*(.+?)\s*$")
FIELD_PATTERN = re.compile(
    rf"^(\s*)({'|'.join(re.escape(x) for x in TARGET_FIELDS)})(\s*:\s*)(.*)$"
)


# =========================
# DOCX ITER HELPERS
# =========================
def iter_block_items(parent):
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(child, parent)
        elif tag == "tbl":
            yield Table(child, parent)


def iter_paragraphs_recursive(parent):
    for block in iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    yield from iter_paragraphs_recursive(cell)


# =========================
# RUN FORMAT HELPERS
# =========================
def get_run_rpr(run):
    """Lấy copy của run properties XML để giữ formatting."""
    rpr = run._r.rPr
    return deepcopy(rpr) if rpr is not None else None


def clear_paragraph_runs(paragraph):
    """Xóa toàn bộ run trong paragraph nhưng giữ paragraph properties."""
    p = paragraph._p
    for child in list(p):
        # chỉ xóa run, giữ nguyên pPr
        if child.tag.endswith("}r"):
            p.remove(child)


def add_run_with_rpr(paragraph, text, rpr=None):
    """Thêm run mới với formatting XML cũ."""
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, deepcopy(rpr))
    return run


def build_char_style_map(paragraph):
    """
    Trả về:
    - full_text: text đầy đủ của paragraph
    - char_styles: list style XML ứng với từng ký tự
    """
    full_text = ""
    char_styles = []

    for run in paragraph.runs:
        t = run.text
        rpr = get_run_rpr(run)
        full_text += t
        char_styles.extend([rpr] * len(t))

    return full_text, char_styles


def normalize_rpr_key(rpr):
    """Tạo key để group các ký tự có cùng style."""
    if rpr is None:
        return None
    return rpr.xml


def rebuild_paragraph_with_styles(paragraph, new_text, new_styles):
    """
    Rebuild paragraph từ text + style map theo từng ký tự.
    Gộp các ký tự liên tiếp có cùng style thành 1 run.
    """
    clear_paragraph_runs(paragraph)

    if not new_text:
        return

    chunks = []
    current_text = new_text[0]
    current_style = new_styles[0]

    for ch, st in zip(new_text[1:], new_styles[1:]):
        if normalize_rpr_key(st) == normalize_rpr_key(current_style):
            current_text += ch
        else:
            chunks.append((current_text, current_style))
            current_text = ch
            current_style = st

    chunks.append((current_text, current_style))

    for text, style in chunks:
        add_run_with_rpr(paragraph, text, style)


# =========================
# CORE REPLACE LOGIC
# =========================
def replace_field_prefix_preserve_format(paragraph, current_subreq):
    """
    Chỉ thay phần prefix:
      Guidance - Purpose:
    thành:
      Guidance - Purpose of 2.1.1:
    và giữ formatting phần còn lại.
    """
    full_text, char_styles = build_char_style_map(paragraph)
    if not full_text.strip():
        return False

    m = FIELD_PATTERN.match(full_text)
    if not m or not current_subreq:
        return False

    indent, field_name, sep, rest = m.groups()

    old_prefix = f"{indent}{field_name}{sep}"
    new_prefix = f"{indent}{field_name} of {current_subreq}{sep}"

    # tránh replace lặp
    if full_text.startswith(new_prefix):
        return False

    old_prefix_len = len(old_prefix)

    # style cho phần prefix mới:
    # ưu tiên lấy style của ký tự đầu tiên trong prefix cũ
    if old_prefix_len > 0 and len(char_styles) >= old_prefix_len:
        prefix_style = char_styles[0]
    elif char_styles:
        prefix_style = char_styles[0]
    else:
        prefix_style = None

    new_prefix_styles = [prefix_style] * len(new_prefix)

    # phần còn lại giữ nguyên style gốc sau old_prefix
    rest_styles = char_styles[old_prefix_len:]

    new_text = new_prefix + full_text[old_prefix_len:]
    new_styles = new_prefix_styles + rest_styles

    # fallback an toàn nếu lệch length
    if len(new_text) != len(new_styles):
        new_styles = [prefix_style] * len(new_text)

    rebuild_paragraph_with_styles(paragraph, new_text, new_styles)
    return True


# =========================
# DOCX PROCESS
# =========================
def process_docx(input_path, output_path):
    doc = Document(input_path)
    current_subreq = None
    updated = 0

    for para in iter_paragraphs_recursive(doc):
        text = para.text.strip()
        if not text:
            continue

        m_sub = SUBREQ_PATTERN.match(text)
        if m_sub:
            current_subreq = m_sub.group(1).strip()
            continue

        if replace_field_prefix_preserve_format(para, current_subreq):
            updated += 1

    doc.save(output_path)
    return updated


# =========================
# BATCH
# =========================
def batch_process():
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    total_files = 0
    total_updates = 0

    for filename in os.listdir(INPUT_FOLDER):
        if not filename.lower().endswith(".docx"):
            continue
        if filename.startswith("~$"):
            continue

        input_path = os.path.join(INPUT_FOLDER, filename)

        base, ext = os.path.splitext(filename)
        output_filename = f"{base}{ext}"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        print(f"Processing: {filename}")
        updated = process_docx(input_path, output_path)
        print(f"  -> Updated {updated} paragraphs")
        print(f"  -> Saved: {output_path}")

        total_files += 1
        total_updates += updated

    print("\n===== DONE =====")
    print(f"Files processed: {total_files}")
    print(f"Total updates: {total_updates}")


if __name__ == "__main__":
    batch_process()