#!/usr/bin/env python3
"""Convert requirement content (PDF/TXT) into a DOCX following a strict template."""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, Iterable, List, Tuple, TypedDict

LABEL_ALIASES: Dict[str, List[str]] = {
    "requirement_id": ["Mã Yêu cầu"],
    "defined_approach_requirements": ["Defined Approach Requirements"],
    "defined_approach_testing_procedures": ["Defined Approach Testing Procedures"],
    "guidance_purpose": ["Guidance - Purpose", "Purpose"],
    "guidance_good_practice": ["Guidance - Good Practice", "Good Practice"],
    "guidance_definitions": ["Guidance - Definitions", "Definitions"],
    "guidance_examples": ["Guidance - Examples", "Examples"],
    "guidance_further_information": ["Guidance - Further Information", "Further Information"],
    "customized_approach_objective": ["Customized Approach Objective"],
    "applicability_notes": ["Applicability Notes"],
}


@dataclass
class RequirementRecord:
    requirement_id: str = "[Nhập ID yêu cầu]"
    defined_approach_requirements: str = "N/A"
    defined_approach_testing_procedures: str = (
        '- "[Nhập ID Test a]": [Chủ thể: Assessor] [Hành động: Examine/Observe/Interview/Review] '
        '[Đối tượng: system/documentation/personnel] để [Hành động chốt chặn: Verify/Confirm] '
        '[Nội dung cần xác minh].\n'
        '- "[Nhập ID Test b]": [Chủ thể: Assessor] [Hành động] [Đối tượng] để [Hành động chốt chặn] '
        '[Nội dung cần xác minh].'
    )
    guidance_purpose: str = "N/A"
    guidance_good_practice: str = "N/A"
    guidance_definitions: str = "N/A"
    guidance_examples: str = "N/A"
    guidance_further_information: str = "N/A"
    customized_approach_objective: str = "N/A"
    applicability_notes: str = "N/A"


def read_input_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            pdfplumber = None

        if pdfplumber:
            # Keep legacy behavior if caller wants a single combined text.
            text_parts: List[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text_parts.append(page_text)
            if text_parts:
                return "\n".join(text_parts)

        pdftotext = shutil.which("pdftotext")
        if not pdftotext:
            raise SystemExit(
                "Không tìm thấy `pdftotext`. Hãy cài poppler-utils hoặc đổi PDF sang TXT trước."
            )
        result = subprocess.run(
            [pdftotext, str(path), "-"], capture_output=True, text=True, check=True
        )
        return result.stdout
    return path.read_text(encoding="utf-8", errors="ignore")


def run_ocrmypdf(input_path: Path) -> Path:
    ocrmypdf = shutil.which("ocrmypdf")
    if not ocrmypdf:
        raise SystemExit("Không tìm thấy `ocrmypdf`. Hãy cài OCRmyPDF trước.")
    tmp_dir = Path(tempfile.mkdtemp(prefix="ocrmypdf_"))
    output_path = tmp_dir / f"{input_path.stem}.ocr.pdf"
    result = subprocess.run(
        [
            ocrmypdf,
            "--deskew",
            "--clean",
            "--optimize",
            "1",
            "--force-ocr",
            str(input_path),
            str(output_path),
        ],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0 or not output_path.exists():
        raise SystemExit(f"OCRmyPDF lỗi:\n{result.stderr.strip()}")
    return output_path


def _normalize_text(text: str) -> str:
    normalized = text.replace("\u00a0", " ")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    return normalized


_INVALID_XML_CHARS_RE = re.compile(
    r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]"
)


def _sanitize_xml_text(value: str) -> str:
    return _INVALID_XML_CHARS_RE.sub("", value)


def _build_heading_pattern(labels: Iterable[str]) -> re.Pattern[str]:
    sorted_labels = sorted(labels, key=len, reverse=True)
    escaped = "|".join(re.escape(label) for label in sorted_labels)
    return re.compile(rf"(?im)(?:^|[ \t]{{2,}})({escaped})\s*:?\s*")


def _extract_sections(normalized: str, labels: Iterable[str]) -> Dict[str, str]:
    pattern = _build_heading_pattern(labels)
    matches = list(pattern.finditer(normalized))
    sections: Dict[str, str] = {}
    if not matches:
        return sections

    for i, match in enumerate(matches):
        label = match.group(1)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(normalized)
        value = normalized[start:end].strip()
        if value:
            sections[label] = value
    return sections


def _extract_requirement_id(normalized: str) -> str | None:
    match = re.search(r"\b\d+\.\d+\.\d+\.\d+\b", normalized)
    if match:
        return match.group(0)
    match = re.search(r"\b\d+\.\d+\.\d+\b", normalized)
    if match:
        return match.group(0)
    return None


def parse_record(text: str) -> RequirementRecord:
    record = RequirementRecord()
    normalized = _normalize_text(text)

    all_labels: List[str] = []
    for aliases in LABEL_ALIASES.values():
        all_labels.extend(aliases)

    sections = _extract_sections(normalized, all_labels)
    if sections:
        for field, aliases in LABEL_ALIASES.items():
            for alias in aliases:
                if alias in sections:
                    setattr(record, field, sections[alias])
                    break

    if record.requirement_id == RequirementRecord.requirement_id:
        inferred = _extract_requirement_id(normalized)
        if inferred:
            record.requirement_id = inferred

    return record


def _cluster_centers(xs: List[float], k: int = 3, iterations: int = 8) -> List[float]:
    if not xs:
        return []
    xs_sorted = sorted(xs)
    if len(xs_sorted) < k:
        return xs_sorted
    centers = [
        xs_sorted[int(len(xs_sorted) * 0.1)],
        xs_sorted[int(len(xs_sorted) * 0.5)],
        xs_sorted[int(len(xs_sorted) * 0.9)],
    ]
    for _ in range(iterations):
        buckets = [[] for _ in range(k)]
        for x in xs_sorted:
            idx = min(range(k), key=lambda i: abs(x - centers[i]))
            buckets[idx].append(x)
        for i in range(k):
            if buckets[i]:
                centers[i] = sum(buckets[i]) / len(buckets[i])
    return centers


def _assign_column(x0: float, boundaries: List[float]) -> int:
    for i, bound in enumerate(boundaries):
        if x0 < bound:
            return i
    return len(boundaries)


def _group_lines(words: List[Dict[str, float]], y_tol: float = 2.0) -> List[List[Dict[str, float]]]:
    words_sorted = sorted(words, key=lambda w: (w["top"], w["x0"]))
    lines: List[List[Dict[str, float]]] = []
    current: List[Dict[str, float]] = []
    current_top: float | None = None
    for word in words_sorted:
        if current_top is None or abs(word["top"] - current_top) <= y_tol:
            current.append(word)
            if current_top is None:
                current_top = word["top"]
        else:
            lines.append(current)
            current = [word]
            current_top = word["top"]
    if current:
        lines.append(current)
    return lines


_FOOTER_PATTERNS = (
    re.compile(r"Payment Card Industry Data Security Standard", re.IGNORECASE),
    re.compile(r"PCI Security Standards Council", re.IGNORECASE),
    re.compile(r"All Rights Reserved", re.IGNORECASE),
    re.compile(r"Page\\s+\\d+", re.IGNORECASE),
    re.compile(r"June\\s+\\d{4}", re.IGNORECASE),
    re.compile(r"Requirements and Testing Procedures", re.IGNORECASE),
    re.compile(r"Guidance", re.IGNORECASE),
    re.compile(r"Reserved\\.?$", re.IGNORECASE),
)


def _lines_to_text(lines: List[List[Dict[str, float]]]) -> str:
    texts: List[str] = []
    for line in lines:
        line_text = " ".join(w["text"] for w in line).strip()
        if not line_text:
            continue
        if any(p.search(line_text) for p in _FOOTER_PATTERNS):
            continue
        if line_text:
            texts.append(line_text)
    return "\n".join(texts)


class ColumnBlock(TypedDict):
    req_id: str
    left: str
    middle: str
    right: str


def _extract_pdf_blocks(path: Path) -> List[ColumnBlock]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise SystemExit(
            "PDF nhiều cột cần `pdfplumber` để tách cột. "
            "Cài: pip install pdfplumber pdfminer.six"
        ) from exc

    blocks: List[ColumnBlock] = []
    current_req_id: str | None = None

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(x_tolerance=1, y_tolerance=3) or []
            if not words:
                continue

            # Filter headers/footers
            filtered = []
            for w in words:
                if w["top"] < 40 or w["bottom"] > page.height - 40:
                    continue
                filtered.append(w)
            if not filtered:
                continue

            xs = [(w["x0"] + w["x1"]) / 2 for w in filtered]
            centers = sorted(_cluster_centers(xs, k=3))
            if len(centers) < 2:
                boundaries: List[float] = []
            else:
                boundaries = [(centers[i] + centers[i + 1]) / 2 for i in range(len(centers) - 1)]

            for w in filtered:
                x_center = (w["x0"] + w["x1"]) / 2
                w["col"] = _assign_column(x_center, boundaries)

            left_words = [w for w in filtered if w["col"] == 0]
            if not left_words:
                continue
            req_map: Dict[str, float] = {}
            left_lines = _group_lines(left_words)
            for line in left_lines:
                first = min(line, key=lambda w: w["x0"])
                match = re.match(r"^(\d+\.\d+\.\d+\.\d+|\d+\.\d+\.\d+)\b", first["text"])
                if not match:
                    continue
                req_id = match.group(1)
                y_top = min(w["top"] for w in line)
                if req_id not in req_map or y_top < req_map[req_id]:
                    req_map[req_id] = y_top
            req_starts: List[Tuple[str, float]] = [(req_id, y) for req_id, y in req_map.items()]

            if not req_starts:
                if current_req_id:
                    lines = _group_lines(filtered)
                    text = _lines_to_text(lines)
                    blocks.append(
                        {
                            "req_id": current_req_id,
                            "left": text,
                            "middle": "",
                            "right": "",
                        }
                    )
                continue

            # Keep all levels; parent requirements should still be extracted when children exist.

            req_starts.sort(key=lambda item: item[1])
            for i, (req_id, y_start_raw) in enumerate(req_starts):
                y_start = max(0, y_start_raw - 8)
                y_end = req_starts[i + 1][1] - 1 if i + 1 < len(req_starts) else page.height - 40
                left_sel = [w for w in filtered if w["col"] == 0 and y_start <= w["top"] < y_end]
                mid_sel = [w for w in filtered if w["col"] == 1 and y_start <= w["top"] < y_end]
                right_sel = [w for w in filtered if w["col"] == 2 and y_start <= w["top"] < y_end]
                right_above = [
                    w
                    for w in filtered
                    if w["col"] == 2 and (y_start - 50) <= w["top"] < y_start
                ]
                left_text = _lines_to_text(_group_lines(left_sel))
                mid_text = _lines_to_text(_group_lines(mid_sel))
                right_lines = _group_lines(right_sel)
                above_lines = _group_lines(right_above)
                if above_lines:
                    filtered_above: List[List[Dict[str, float]]] = []
                    for line in above_lines:
                        line_text = " ".join(w["text"] for w in line).strip()
                        if not line_text:
                            continue
                        if f"Requirement {req_id}" in line_text or line_text == "Purpose":
                            filtered_above.append(line)
                    if filtered_above:
                        right_lines = filtered_above + right_lines
                right_text = _lines_to_text(right_lines)
                blocks.append(
                    {
                        "req_id": req_id,
                        "left": left_text,
                        "middle": mid_text,
                        "right": right_text,
                    }
                )
                current_req_id = req_id

    return blocks


def _split_text_blocks_by_requirement(text: str) -> List[ColumnBlock]:
    normalized = _normalize_text(text)
    pattern = re.compile(r"(?m)^(?P<id>\d+\.\d+\.\d+\.\d+|\d+\.\d+\.\d+)\b")
    matches = list(pattern.finditer(normalized))
    if not matches:
        return []
    blocks: List[ColumnBlock] = []
    for i, match in enumerate(matches):
        req_id = match.group("id")
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(normalized)
        chunk = normalized[start:end].strip()
        if chunk:
            blocks.append(
                {
                    "req_id": req_id,
                    "left": chunk,
                    "middle": "",
                    "right": "",
                }
            )
    return blocks


def _split_by_heading(text: str, headings: List[str]) -> Dict[str, str]:
    if not text:
        return {}
    pattern = _build_heading_pattern(headings)
    matches = list(pattern.finditer(text))
    if not matches:
        return {}
    sections: Dict[str, str] = {}
    for i, match in enumerate(matches):
        label = match.group(1)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        value = text[start:end].strip()
        if value:
            sections[label] = value
    return sections


def _split_by_heading_lines(text: str, headings: List[str]) -> Dict[str, str]:
    if not text:
        return {}
    escaped = "|".join(re.escape(h) for h in sorted(headings, key=len, reverse=True))
    pattern = re.compile(rf"(?im)^\s*({escaped})\s*:?\s*$")
    matches = list(pattern.finditer(text))
    if not matches:
        return {}
    sections: Dict[str, str] = {}
    for i, match in enumerate(matches):
        label = match.group(1)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        value = text[start:end].strip()
        if value:
            sections[label] = value
    return sections


def _strip_leading_req_id(text: str, req_id: str) -> str:
    if not text:
        return text
    pattern = re.compile(rf"^\\s*{re.escape(req_id)}\\b\\s*", re.MULTILINE)
    return pattern.sub("", text, count=1).lstrip()


def _remove_inline_req_id(text: str, req_id: str) -> str:
    if not text:
        return text
    pattern = re.compile(rf"(?i)(?<!Requirement\s)(?<!\S){re.escape(req_id)}(?!\S)")
    return pattern.sub("", text)


def _cleanup_column_text(text: str) -> str:
    if not text:
        return text
    text = re.sub(r"^.*(June\s+\d{4}|Page\s+\d+).*$", "", text, flags=re.IGNORECASE | re.MULTILINE)
    text = re.sub(r"Requirements and Testing Procedures", "", text, flags=re.IGNORECASE)
    text = re.sub(r"All Rights Reserved", "", text, flags=re.IGNORECASE)
    text = re.sub(r"PCI Security Standards Council", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\n\s*\n", "\n", text)
    return text.strip()


def _unwrap_lines(text: str) -> str:
    if not text:
        return text
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    merged: List[str] = []
    for line in lines:
        is_bullet = line.startswith("•") or line.startswith("- ")
        if is_bullet:
            merged.append(line)
            continue
        if merged and merged[-1].startswith(("•", "- ")):
            # Continuation of a bullet line; append unless it looks like a new paragraph.
            if merged[-1].endswith((".", "?", "!")) and line and line[0].isupper():
                merged.append(line)
            else:
                merged[-1] = f"{merged[-1]} {line}"
        elif merged:
            if merged[-1].endswith((".", "?", "!")) and line and line[0].islower():
                line = line[0].upper() + line[1:]
            merged[-1] = f"{merged[-1]} {line}"
        else:
            merged.append(line)
    return "\n".join(merged)


def _normalize_procedure_ids(text: str) -> str:
    if not text:
        return text
    # Ensure each .a/.b/.c procedure starts on its own line.
    text = re.sub(r"\s+(\d+\.\d+\.\d+\.[a-z])\s+", r"\n\1 ", text, flags=re.IGNORECASE)
    lines = text.splitlines()
    normalized: List[str] = []
    pattern = re.compile(r"\b(\d+\.\d+\.\d+\.[a-z])\b", re.IGNORECASE)
    for line in lines:
        match = pattern.search(line)
        if match and match.start() > 0:
            proc_id = match.group(1)
            before = line[:match.start()].strip()
            after = line[match.end():].strip()
            remainder = " ".join(part for part in [before, after] if part)
            line = f"{proc_id} {remainder}".strip()
        line = re.sub(r"\s{2,}", " ", line)
        normalized.append(line)
    return "\n".join(normalized)


def _extract_procedure_items(text: str) -> str:
    if not text:
        return text
    items: List[str] = []
    pattern = re.compile(r"\b(\d+\.\d+\.\d+\.[a-z])\b", re.IGNORECASE)
    matches = list(pattern.finditer(text))
    for i, match in enumerate(matches):
        proc_id = match.group(1)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        chunk = text[start:end].strip()
        if chunk.lower().startswith(proc_id.lower()):
            chunk = chunk[len(proc_id):].strip()
        chunk = _unwrap_lines(chunk)
        if chunk:
            items.append(f"{proc_id} {chunk}".strip())
    if items:
        return "\n".join(items)
    return _normalize_procedure_ids(_unwrap_lines(text))


def _remove_procedure_fragments(text: str) -> str:
    if not text:
        return text
    cleaned: List[str] = []
    line_pattern = re.compile(r"^\s*\d+\.\d+\.\d+\.[a-z]\b", re.IGNORECASE)
    token_pattern = re.compile(r"\b(\d+\.\d+\.\d+\.[a-z])\b", re.IGNORECASE)
    for line in text.splitlines():
        if line_pattern.match(line):
            continue
        line = token_pattern.sub("", line).strip()
        if not line:
            continue
        line = re.sub(r"\s{2,}", " ", line)
        cleaned.append(line)
    return "\n".join(cleaned)


def parse_record_from_columns(left: str, middle: str, right: str, req_id: str) -> RequirementRecord:
    record = RequirementRecord()
    record.requirement_id = req_id

    left = _cleanup_column_text(
        _remove_inline_req_id(_strip_leading_req_id(_normalize_text(left), req_id), req_id)
    )
    middle = _cleanup_column_text(
        _remove_inline_req_id(_strip_leading_req_id(_normalize_text(middle), req_id), req_id)
    )
    right = _cleanup_column_text(
        _remove_inline_req_id(_strip_leading_req_id(_normalize_text(right), req_id), req_id)
    )

    left = re.sub(
        r"\s*(Customized Approach Objective|Applicability Notes)\b",
        r"\n\1\n",
        left,
        flags=re.IGNORECASE,
    )
    left_sections = _split_by_heading_lines(left, ["Customized Approach Objective", "Applicability Notes"])
    if left_sections:
        # Text before first heading is Defined Approach Requirements.
        first_heading = next(iter(left_sections.keys()))
        first_index = left.find(first_heading)
        if first_index > 0:
            record.defined_approach_requirements = _remove_procedure_fragments(
                _unwrap_lines(left[:first_index].strip())
            )
        record.customized_approach_objective = _unwrap_lines(
            left_sections.get("Customized Approach Objective", record.customized_approach_objective)
        )
        record.applicability_notes = _unwrap_lines(
            left_sections.get("Applicability Notes", record.applicability_notes)
        )
    elif left.strip():
        record.defined_approach_requirements = _remove_procedure_fragments(
            _unwrap_lines(left.strip())
        )

    if middle.strip():
        record.defined_approach_testing_procedures = _extract_procedure_items(middle.strip())

    right_sections = _split_by_heading(
        right,
        ["Purpose", "Good Practice", "Definitions", "Examples", "Further Information"],
    )
    if right_sections:
        if "Purpose" in right_sections:
            record.guidance_purpose = _unwrap_lines(
                right_sections.get("Purpose", record.guidance_purpose)
            )
        else:
            first_heading = next(iter(right_sections.keys()))
            first_index = right.find(first_heading)
            if first_index > 0:
                record.guidance_purpose = _unwrap_lines(right[:first_index].strip())
        record.guidance_good_practice = _unwrap_lines(
            right_sections.get("Good Practice", record.guidance_good_practice)
        )
        record.guidance_definitions = _unwrap_lines(
            right_sections.get("Definitions", record.guidance_definitions)
        )
        record.guidance_examples = _unwrap_lines(
            right_sections.get("Examples", record.guidance_examples)
        )
        record.guidance_further_information = _unwrap_lines(
            right_sections.get("Further Information", record.guidance_further_information)
        )
    elif right.strip():
        # If headings are missing, keep all guidance in Purpose as a fallback.
        record.guidance_purpose = _unwrap_lines(right.strip())

    return record




def write_docx(record: RequirementRecord, output_file: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit(
            "Thiếu thư viện `python-docx`. Chạy: pip install python-docx"
        ) from exc

    doc = Document()
    data = asdict(record)

    def add_paragraph_safe(text: str) -> None:
        doc.add_paragraph(_sanitize_xml_text(text))

    def add_label_value(label: str, value: str) -> None:
        if value.strip().upper() == "N/A":
            return
        p = doc.add_paragraph()
        run_label = p.add_run(_sanitize_xml_text(label))
        run_label.bold = True
        p.add_run(_sanitize_xml_text(value))

    add_label_value('Mã Yêu cầu: ', f'"{data["requirement_id"]}"')
    add_label_value(
        "Defined Approach Requirements: ", data["defined_approach_requirements"]
    )

    dap_lines = [
        line.strip()
        for line in data["defined_approach_testing_procedures"].splitlines()
        if line.strip() and line.strip().upper() != "N/A"
    ]
    if dap_lines:
        if not dap_lines[0].startswith(("•", "- ")):
            add_label_value("Defined Approach Testing Procedures: ", dap_lines[0])
            rest_lines = dap_lines[1:]
        else:
            add_label_value("Defined Approach Testing Procedures: ", "")
            rest_lines = dap_lines
        for line in rest_lines:
            add_paragraph_safe(line)

    add_label_value("Guidance - Purpose: ", data["guidance_purpose"])
    add_label_value("Guidance - Good Practice: ", data["guidance_good_practice"])
    add_label_value("Guidance - Definitions: ", data["guidance_definitions"])
    add_label_value("Guidance - Examples: ", data["guidance_examples"])
    add_label_value(
        "Guidance - Further Information: ", data["guidance_further_information"]
    )
    add_label_value(
        "Customized Approach Objective: ", data["customized_approach_objective"]
    )
    add_label_value("Applicability Notes: ", data["applicability_notes"])

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Chuyển PDF/TXT sang DOCX theo format chuẩn yêu cầu."
    )
    parser.add_argument("input", type=Path, help="Đường dẫn file .pdf hoặc .txt")
    parser.add_argument(
        "-o", "--output", type=Path, default=Path("output/converted.docx"), help="File DOCX đầu ra"
    )
    parser.add_argument(
        "--id",
        type=str,
        default="",
        help="Chỉ xuất 1 requirement cụ thể (ví dụ: 1.2.7).",
    )
    args = parser.parse_args()

    if args.input.suffix.lower() == ".pdf":
        input_path = run_ocrmypdf(args.input)
        blocks = _extract_pdf_blocks(input_path)
        if not blocks:
            text = read_input_text(input_path)
            blocks = _split_text_blocks_by_requirement(text)
        if not blocks:
            text = read_input_text(input_path)
            record = parse_record(text)
            write_docx(record, args.output)
            print(f"Đã tạo: {args.output}")
            return

        # Merge blocks by requirement id in order of appearance
        merged: Dict[str, List[ColumnBlock]] = {}
        order: List[str] = []
        for block in blocks:
            req_id = block["req_id"]
            if req_id not in merged:
                merged[req_id] = []
                order.append(req_id)
            merged[req_id].append(block)

        if args.id:
            target = args.id.strip()
            if target not in merged:
                print(f"Không tìm thấy requirement: {target}")
                print(f"Các requirement phát hiện: {', '.join(order)}")
                return
            blocks_for_id = merged[target]
            left = "\n".join(b["left"] for b in blocks_for_id if b["left"])
            middle = "\n".join(b["middle"] for b in blocks_for_id if b["middle"])
            right = "\n".join(b["right"] for b in blocks_for_id if b["right"])
            record = parse_record_from_columns(left, middle, right, target)
            write_docx(record, args.output)
            print(f"Đã tạo: {args.output} (requirement {target})")
            return

        if len(order) == 1 and not args.id:
            blocks_for_id = merged[order[0]]
            left = "\n".join(b["left"] for b in blocks_for_id if b["left"])
            middle = "\n".join(b["middle"] for b in blocks_for_id if b["middle"])
            right = "\n".join(b["right"] for b in blocks_for_id if b["right"])
            record = parse_record_from_columns(left, middle, right, order[0])
            write_docx(record, args.output)
            print(f"Đã tạo: {args.output}")
            return

        output_dir = args.output if args.output.is_dir() else args.output.parent
        for req_id in order:
            blocks_for_id = merged[req_id]
            left = "\n".join(b["left"] for b in blocks_for_id if b["left"])
            middle = "\n".join(b["middle"] for b in blocks_for_id if b["middle"])
            right = "\n".join(b["right"] for b in blocks_for_id if b["right"])
            record = parse_record_from_columns(left, middle, right, req_id)
            safe_id = re.sub(r"[^0-9A-Za-z._-]+", "_", req_id)
            out_path = output_dir / f"{safe_id}.docx"
            write_docx(record, out_path)
        print(f"Đã tạo {len(order)} file trong: {output_dir}")
        return

    text = read_input_text(args.input)
    record = parse_record(text)
    write_docx(record, args.output)
    print(f"Đã tạo: {args.output}")


if __name__ == "__main__":
    main()
