#!/usr/bin/env python3
import argparse
import csv
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from docx import Document

REQ_CODE_RE = re.compile(r"\b\d+(?:\.\d+){2,}\b(?!\.[a-z])")
REQ_CODE_FROM_TEST_RE = re.compile(r"\b(\d+(?:\.\d+){2,})\.[a-z]\b")
TEST_CODE_RE = re.compile(r"\b(\d+(?:\.\d+){2,}\.[a-z])\b")
REQ_CODE_START_RE = re.compile(r"^\s*(\d+(?:\.\d+){2,})\b(?!\.[a-z])")
TEST_CODE_START_RE = re.compile(r"^\s*(\d+(?:\.\d+){2,}\.[a-z])\b")

MARKERS = [
    "Defined Approach Requirements",
    "Customized Approach Objective",
    "Applicability Notes",
    "Defined Approach Testing Procedures",
    "Purpose",
    "Good Practice",
    "Guidance - Purpose",
    "Guidance - Good Practice",
    "Guidance - Definitions",
    "Guidance - Examples",
    "Guidance - Further Information",
    "Definitions",
    "Examples",
    "Further Information",
]

MARKER_RE = re.compile("|".join(re.escape(m) for m in MARKERS), re.IGNORECASE)
C0_MARKERS = [
    "Defined Approach Requirements",
    "Customized Approach Objective",
    "Applicability Notes",
]
C2_MARKERS = [
    "Purpose",
    "Good Practice",
    "Definitions",
    "Examples",
    "Further Information",
    "Guidance - Purpose",
    "Guidance - Good Practice",
    "Guidance - Definitions",
    "Guidance - Examples",
    "Guidance - Further Information",
]


@dataclass
class RequirementItem:
    requirements: str = ""
    objective: str = ""
    applicability: str = ""
    purpose: str = ""
    good_practice: str = ""
    definitions: str = ""
    examples: str = ""
    further_information: str = ""
    tests: List[Tuple[str, str]] = field(default_factory=list)


def normalize_space(text: str) -> str:
    text = re.sub(r"\(\s*continued\s+on\s+next\s+page\s*\)", "", text, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", text).strip()


def find_req_code(text: str) -> str | None:
    m = REQ_CODE_FROM_TEST_RE.search(text or "")
    if m:
        return m.group(1)
    m = REQ_CODE_RE.search(text or "")
    if m:
        return m.group(0)
    return None


def find_req_code_at_start(text: str) -> str | None:
    m = TEST_CODE_START_RE.search(text or "")
    if m:
        return m.group(1).rsplit(".", 1)[0]
    m = REQ_CODE_START_RE.search(text or "")
    if m:
        return m.group(1)
    return None


def find_test_code_at_start(text: str) -> str | None:
    m = TEST_CODE_START_RE.search(text or "")
    if m:
        return m.group(1)
    return None


def extract_section(text: str, marker: str) -> str:
    if not text:
        return ""
    def _search(t: str) -> str:
        pattern = re.compile(
            re.escape(marker) + r"\s*(.*?)\s*(?=(" + "|".join(re.escape(m) for m in MARKERS) + r")|$)",
            re.IGNORECASE | re.DOTALL,
        )
        m = pattern.search(t)
        if not m:
            return ""
        return normalize_space(m.group(1))

    out = _search(text)
    if out:
        return out

    # Handle duplicated marker like "Examples Examples ..."
    dup_pat = re.compile(rf"({re.escape(marker)})\s+{re.escape(marker)}", re.IGNORECASE)
    if dup_pat.search(text):
        text2 = dup_pat.sub(r"\1", text, count=1)
        out2 = _search(text2)
        if not out2:
            return ""
        # keep a single leading marker as part of content (e.g., "Examples of ...")
        if out2.lower().startswith(marker.lower()):
            return out2
        return f"{marker} {out2}"
    return ""


def split_c0_sections(text: str) -> Dict[str, str]:
    if not text:
        return {}
    headers: List[Tuple[int, int, str]] = []
    def _is_boundary(t: str, idx: int) -> bool:
        if idx <= 0:
            return True
        prev = t[idx - 1]
        if prev in ".\n":
            return True
        if idx >= 2 and t[idx - 2:idx] in {". ", ".\n"}:
            return True
        return False
    for m in re.finditer(
        "|".join(re.escape(x) for x in C0_MARKERS),
        text,
        flags=re.IGNORECASE,
    ):
        if not _is_boundary(text, m.start()):
            continue
        prefix = text[max(0, m.start() - 20) : m.start()].lower()
        if "refer to" in prefix or "see " in prefix:
            continue
        headers.append((m.start(), m.end(), m.group(0)))
    if not headers:
        return {}
    headers.sort(key=lambda x: x[0])
    out: Dict[str, str] = {}
    for i, (start, end, name) in enumerate(headers):
        next_start = headers[i + 1][0] if i + 1 < len(headers) else len(text)
        chunk = normalize_space(text[end:next_start])
        if chunk:
            key = next((m for m in C0_MARKERS if m.lower() == name.lower()), name)
            out[key] = chunk
    return out


def split_c2_sections(text: str) -> Dict[str, str]:
    if not text:
        return {}
    headers: List[Tuple[int, int, str]] = []

    def _is_boundary(t: str, idx: int) -> bool:
        if idx <= 0:
            return True
        prev = t[idx - 1]
        if prev in ".!?:\n":
            return True
        if idx >= 2 and t[idx - 2:idx] in {
            ". ",
            ".\n",
            "! ",
            "!\n",
            "? ",
            "?\n",
            ": ",
            ":\n",
        }:
            return True
        return False

    marker_pat = re.compile(
        r"(?<!\w)(" + "|".join(re.escape(x) for x in C2_MARKERS) + r")(?!\w)",
        flags=re.IGNORECASE,
    )
    for m in marker_pat.finditer(text):
        if not _is_boundary(text, m.start()):
            continue
        prefix = text[max(0, m.start() - 20) : m.start()].lower()
        if "refer to" in prefix or "see " in prefix:
            continue
        headers.append((m.start(), m.end(), m.group(1)))
    if not headers:
        return {}
    headers.sort(key=lambda x: x[0])
    out: Dict[str, str] = {}
    if headers[0][0] > 0:
        leading = normalize_space(text[: headers[0][0]])
        if leading:
            out["Purpose"] = leading
    for i, (start, end, name) in enumerate(headers):
        next_start = headers[i + 1][0] if i + 1 < len(headers) else len(text)
        chunk = normalize_space(text[end:next_start])
        if chunk:
            key = next((m for m in C2_MARKERS if m.lower() == name.lower()), name)
            out[key] = chunk
    return out


def _looks_like_testing(text: str, req_code: str) -> bool:
    if not text:
        return False
    if re.search(r"defined approach testing procedures", text, re.IGNORECASE):
        return True
    if TEST_CODE_RE.search(text):
        return True
    if req_code and re.match(rf"^\s*{re.escape(req_code)}\.\w", text):
        return True
    return False


def strip_leading_code(text: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+){2,}\b\s*", "", text).strip()


def extract_requirement_fallback(parts: Iterable[str], req_code: str) -> str:
    if not req_code:
        return ""
    for p in parts:
        if not p:
            continue
        lower = p.strip().lower()
        if lower.startswith("defined approach requirements"):
            continue
        if lower.startswith("defined approach testing procedures"):
            continue
        if re.match(rf"^\s*{re.escape(req_code)}\b", p):
            return strip_leading_code(p)
    return ""


def split_tests(text: str) -> List[Tuple[str, str]]:
    tests = []
    if not text:
        return tests
    for m in TEST_CODE_RE.finditer(text):
        code = m.group(1)
        start = m.end()
        end = len(text)
        nxt = TEST_CODE_RE.search(text, start)
        if nxt:
            end = nxt.start()
        desc = normalize_space(text[start:end])
        mm = MARKER_RE.search(desc)
        if mm:
            desc = normalize_space(desc[: mm.start()])
        tests.append((code, desc))
    return tests


def format_test_desc(desc: str) -> str:
    if not desc:
        return ""
    desc = desc.strip()
    # remove leading numbering artifacts like "1.", "1.1", or "1. 1.1"
    desc = re.sub(r"^\s*(?:\d+(?:\.\d+)*\.?\s*)+", "", desc)
    return desc


def next_letter(used: set[str]) -> str:
    for c in "abcdefghijklmnopqrstuvwxyz":
        if c not in used:
            return c
    return "a"


def is_marker_only_row(combined: str, c0: str, c1: str) -> str | None:
    lower = combined.strip().lower()
    for m in MARKERS:
        if lower == m.lower() or c0.lower() == m.lower() or c1.lower() == m.lower():
            return m
    return None


def apply_pending_marker(item: RequirementItem, pending: str, candidate: str) -> bool:
    if pending.lower().startswith("guidance - "):
        pending = pending[11:]
    if pending == "Defined Approach Testing Procedures":
        return False

    if not candidate:
        return False
    if candidate.strip().lower() in (m.lower() for m in MARKERS):
        return False

    if pending == "Customized Approach Objective" and not item.objective:
        item.objective = candidate
        return True
    if pending == "Applicability Notes" and not item.applicability:
        item.applicability = candidate
        return True
    if pending == "Purpose" and not item.purpose:
        item.purpose = candidate
        return True
    if pending == "Good Practice" and not item.good_practice:
        item.good_practice = candidate
        return True
    if pending == "Definitions" and not item.definitions:
        item.definitions = candidate
        return True
    if pending == "Examples" and not item.examples:
        item.examples = candidate
        return True
    if pending == "Further Information" and not item.further_information:
        item.further_information = candidate
        return True

    return False


def update_sections(
    item: RequirementItem,
    c0: str,
    c2: str,
    parts_c0: List[str],
    req_code: str,
    last_section_by_req: Dict[str, str] | None = None,
) -> None:
    # Fixed-column parsing:
    # c0 -> requirements/objective/applicability
    # c2 -> guidance (purpose/good practice/definitions/examples/further info)
    c0_sections = split_c0_sections(c0)
    c2_sections = split_c2_sections(c2)

    req_text = c0_sections.get("Defined Approach Requirements") or extract_section(
        c0, "Defined Approach Requirements"
    )
    if req_text and not item.requirements:
        item.requirements = strip_leading_code(req_text)
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "requirements"
    if not item.requirements:
        fallback_req = extract_requirement_fallback(parts_c0, req_code)
        if fallback_req:
            item.requirements = fallback_req
            if last_section_by_req is not None:
                last_section_by_req[req_code] = "requirements"

    obj_text = c0_sections.get("Customized Approach Objective") or extract_section(
        c0, "Customized Approach Objective"
    )
    if obj_text and not item.objective:
        item.objective = obj_text
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "objective"
    if c0.lower().startswith("customized approach objective"):
        c0_obj = extract_section(c0, "Customized Approach Objective")
        if not c0_obj:
            c0_obj = re.sub(
                r"^customized approach objective\\s*",
                "",
                c0,
                flags=re.IGNORECASE,
            ).strip()
        if c0_obj:
            item.objective = c0_obj
            if last_section_by_req is not None:
                last_section_by_req[req_code] = "objective"

    app_text = c0_sections.get("Applicability Notes")
    if app_text and not item.applicability:
        item.applicability = app_text
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "applicability"

    purpose_text = c2_sections.get("Purpose") or c2_sections.get("Guidance - Purpose")
    if purpose_text and not item.purpose:
        item.purpose = purpose_text
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "purpose"

    gp_text = c2_sections.get("Good Practice") or c2_sections.get("Guidance - Good Practice")
    if gp_text and not item.good_practice:
        item.good_practice = gp_text
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "good_practice"

    def_text = c2_sections.get("Definitions") or c2_sections.get("Guidance - Definitions")
    if def_text and not item.definitions:
        item.definitions = def_text
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "definitions"

    ex_text = c2_sections.get("Examples") or c2_sections.get("Guidance - Examples")
    if ex_text and not item.examples:
        item.examples = ex_text
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "examples"

    fi_text = c2_sections.get("Further Information") or c2_sections.get(
        "Guidance - Further Information"
    )
    if fi_text and not item.further_information:
        item.further_information = fi_text
        if last_section_by_req is not None:
            last_section_by_req[req_code] = "further_information"


def update_tests(
    item: RequirementItem,
    req_code: str,
    c1: str,
    pending_marker_by_req: Dict[str, Tuple[str, str]],
) -> None:
    test_block = extract_section(c1, "Defined Approach Testing Procedures")
    if test_block:
        tests = split_tests(test_block)
        if not tests:
            if re.match(rf"^\s*{re.escape(req_code)}\b", test_block):
                tests = [(req_code, strip_leading_code(test_block))]
            else:
                used_letters = {t[0].split(".")[-1] for t in item.tests if "." in t[0]}
                letter = next_letter(used_letters)
                tests = [(f"{req_code}.{letter}", strip_leading_code(test_block))]
        item.tests.extend(tests)
        pending_marker_by_req.pop(req_code, None)
        return

    tests = split_tests(c1)
    if tests:
        item.tests.extend(tests)
        pending_marker_by_req.pop(req_code, None)
        return
    if c1 and re.match(rf"^\s*{re.escape(req_code)}\b", c1):
        item.tests.append((req_code, strip_leading_code(c1)))
        pending_marker_by_req.pop(req_code, None)
        return

    pending_req = pending_marker_by_req.get(req_code)
    if pending_req and pending_req[0] == "Defined Approach Testing Procedures" and c1:
        tests = split_tests(c1)
        if tests:
            item.tests.extend(tests)
            pending_marker_by_req.pop(req_code, None)
            return
        if c1.startswith(req_code):
            used_letters = {t[0].split(".")[-1] for t in item.tests if "." in t[0]}
            letter = next_letter(used_letters)
            item.tests.append((f"{req_code}.{letter}", strip_leading_code(c1)))
            pending_marker_by_req.pop(req_code, None)


def dedup_tests(item: RequirementItem) -> None:
    seen = set()
    uniq = []
    for code, desc in item.tests:
        if code in seen:
            continue
        seen.add(code)
        uniq.append((code, desc))
    item.tests = uniq


def _append_test_continuation(item: RequirementItem, text: str) -> bool:
    if not item.tests or not text:
        return False
    # avoid appending if the line looks like a new marker or a new test code
    if MARKER_RE.search(text) or TEST_CODE_RE.search(text):
        return False
    code, desc = item.tests[-1]
    joined = normalize_space(f"{desc} {text}")
    item.tests[-1] = (code, joined)
    return True


def _clean_continuation(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^\(?continued\)?\s*", "", text, flags=re.IGNORECASE)
    return text.strip()


def _is_continued_header(text: str) -> bool:
    if not text:
        return False
    return bool(re.search(r"\bcontinued\b", text, flags=re.IGNORECASE))


def _append_section_continuation(item: RequirementItem, section: str, text: str) -> bool:
    if not text:
        return False
    if MARKER_RE.search(text) or TEST_CODE_RE.search(text):
        return False
    cleaned = _clean_continuation(text)
    if not cleaned:
        return False
    if section == "requirements" and item.requirements:
        item.requirements = normalize_space(f"{item.requirements} {cleaned}")
        return True
    if section == "objective" and item.objective:
        item.objective = normalize_space(f"{item.objective} {cleaned}")
        return True
    if section == "applicability" and item.applicability:
        item.applicability = normalize_space(f"{item.applicability} {cleaned}")
        return True
    if section == "purpose" and item.purpose:
        item.purpose = normalize_space(f"{item.purpose} {cleaned}")
        return True
    if section == "good_practice" and item.good_practice:
        item.good_practice = normalize_space(f"{item.good_practice} {cleaned}")
        return True
    if section == "definitions" and item.definitions:
        item.definitions = normalize_space(f"{item.definitions} {cleaned}")
        return True
    if section == "examples" and item.examples:
        item.examples = normalize_space(f"{item.examples} {cleaned}")
        return True
    if section == "further_information" and item.further_information:
        item.further_information = normalize_space(f"{item.further_information} {cleaned}")
        return True
    return False


def _bullet_newlines(text: str) -> str:
    if "•" not in text:
        return text
    text = re.sub(r"\s*•\s*", " • ", text).strip()
    text = text.replace(" • ", "\n• ")
    return text


def format_requirement(code: str, item: RequirementItem) -> List[str]:
    lines = _code_lines(code)
    if item.requirements:
        lines.append(f"Defined Approach Requirements: {_bullet_newlines(item.requirements)}")
    if item.tests:
        if len(item.tests) == 1 and item.tests[0][0] == code:
            lines.append(
                f"Defined Approach Testing Procedures: {format_test_desc(item.tests[0][1])}"
            )
        else:
            lines.append("Defined Approach Testing Procedures:")
            for tcode, desc in item.tests:
                lines.append(f'- "{tcode}": {format_test_desc(desc)}')
    if item.objective:
        lines.append(f"Customized Approach Objective: {_bullet_newlines(item.objective)}")
    if item.applicability:
        lines.append(f"Applicability Notes: {_bullet_newlines(item.applicability)}")
    if item.purpose:
        lines.append(f"Guidance - Purpose: {_bullet_newlines(item.purpose)}")
    if item.good_practice:
        lines.append(f"Guidance - Good Practice: {_bullet_newlines(item.good_practice)}")
    if item.definitions:
        lines.append(f"Guidance - Definitions: {_bullet_newlines(item.definitions)}")
    if item.examples:
        lines.append(f"Guidance - Examples: {_bullet_newlines(item.examples)}")
    if item.further_information:
        lines.append(f"Guidance - Further Information: {_bullet_newlines(item.further_information)}")
    lines.append("---")
    return lines


def parse_csv(input_path: Path) -> Dict[str, RequirementItem]:
    data: Dict[str, RequirementItem] = {}
    last_req_by_table: Dict[str, str] = {}
    pending_marker_by_table: Dict[str, Tuple[str, str]] = {}
    pending_marker_by_req: Dict[str, Tuple[str, str]] = {}
    last_section_by_req: Dict[str, str] = {}
    last_req_global: str | None = None

    with input_path.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            table_idx = row.get("__table_index__", "")
            c0 = row.get("c0", "")
            c1 = row.get("c1", "")
            c2 = row.get("c2", "")
            ocr_row = table_idx == "0"
            explicit_req_code = False
            if ocr_row:
                req_code = (
                    find_req_code_at_start(c0)
                    or find_req_code_at_start(c1)
                    or find_req_code_at_start(c2)
                )
                if req_code:
                    explicit_req_code = True
                if not req_code:
                    test_code = (
                        find_test_code_at_start(c0)
                        or find_test_code_at_start(c1)
                        or find_test_code_at_start(c2)
                    )
                    if test_code and "." in test_code:
                        req_code = test_code.rsplit(".", 1)[0]
            else:
                req_code = (
                    find_req_code_at_start(c0)
                    or find_req_code_at_start(c1)
                    or find_req_code_at_start(c2)
                )
                if req_code:
                    explicit_req_code = True
                if not req_code:
                    c0_lower = (c0 or "").strip().lower()
                    c1_lower = (c1 or "").strip().lower()
                    if c0_lower.startswith("defined approach requirements") or c0_lower.startswith(
                        "defined approach testing procedures"
                    ):
                        req_code = find_req_code(c0)
                        if req_code:
                            explicit_req_code = True
                    elif c1_lower.startswith("defined approach testing procedures"):
                        req_code = find_req_code(c1)
                        if req_code:
                            explicit_req_code = True
                req_code = req_code or last_req_by_table.get(table_idx) or last_req_global
            if not req_code or req_code.count(".") < 2:
                # Continuation row without a detectable requirement code:
                # try to append to the last known section or test description.
                c0s = (c0 or "").strip()
                c1s = (c1 or "").strip()
                c2s = (c2 or "").strip()
                last_req = last_req_by_table.get(table_idx) or last_req_global
                if last_req:
                    item = data.get(last_req)
                    if item:
                        section = last_section_by_req.get(last_req)
                        if section and (c0s or c2s):
                            cont = c0s or c2s
                            if _append_section_continuation(item, section, cont):
                                continue
                        if c1s and _append_test_continuation(item, c1s):
                            continue
                continue

            item = data.setdefault(req_code, RequirementItem())
            last_req_by_table[table_idx] = req_code
            last_req_global = req_code

            c0s = (c0 or "").strip()
            c1s = (c1 or "").strip()
            c2s = (c2 or "").strip()
            c0_has_req = bool(find_req_code(c0s) or find_req_code_at_start(c0s))
            c2_has_req = bool(find_req_code(c2s) or find_req_code_at_start(c2s))

            # Apply any pending marker first (so continuation text doesn't swallow it).
            pending = pending_marker_by_table.get(table_idx)
            if pending:
                pending_marker, source_col = pending
                candidate = c0s if source_col == "c0" else c1s if source_col == "c1" else c2s
                if apply_pending_marker(item, pending_marker, candidate):
                    last_section_by_req[req_code] = {
                        "Customized Approach Objective": "objective",
                        "Applicability Notes": "applicability",
                        "Purpose": "purpose",
                        "Good Practice": "good_practice",
                        "Definitions": "definitions",
                        "Examples": "examples",
                        "Further Information": "further_information",
                    }.get(pending_marker, last_section_by_req.get(req_code, ""))
                    if source_col == "c0":
                        c0s = ""
                    elif source_col == "c1":
                        c1s = ""
                    else:
                        c2s = ""
                    pending_marker_by_table.pop(table_idx, None)
                    pending_marker_by_req.pop(req_code, None)

            # If this row doesn't introduce a new requirement code and c1 looks like
            # a continuation of the previous test, append it and clear c1s.
            if not explicit_req_code and c1s and not TEST_CODE_RE.search(c1s):
                if _append_test_continuation(item, c1s):
                    c1s = ""
            # If this row doesn't introduce a new requirement code and c0/c2 looks like
            # a continuation of the previous section, append it and clear.
            if not explicit_req_code or (section := last_section_by_req.get(req_code)):
                if section:
                    if c0s and not c0_has_req:
                        mm = MARKER_RE.search(c0s)
                        if mm:
                            marker = mm.group(0)
                            prefix = c0s[: mm.start()].strip()
                            suffix = c0s[mm.end() :].strip()
                            if prefix:
                                if item.requirements:
                                    _append_section_continuation(item, "requirements", prefix)
                                else:
                                    _append_section_continuation(item, section, prefix)
                            if suffix:
                                if apply_pending_marker(item, marker, suffix):
                                    last_section_by_req[req_code] = {
                                        "Customized Approach Objective": "objective",
                                        "Applicability Notes": "applicability",
                                        "Purpose": "purpose",
                                        "Good Practice": "good_practice",
                                        "Definitions": "definitions",
                                        "Examples": "examples",
                                        "Further Information": "further_information",
                                    }.get(marker, last_section_by_req.get(req_code, ""))
                            else:
                                pending_marker_by_table[table_idx] = (marker, "c0")
                                pending_marker_by_req[req_code] = (marker, "c0")
                            c0s = ""
                        elif item.requirements and _append_section_continuation(item, "requirements", c0s):
                            c0s = ""
                        elif _append_section_continuation(item, section, c0s):
                            c0s = ""
                    if c2s and not c2_has_req and _append_section_continuation(item, section, c2s):
                        c2s = ""
            guidance_text = c2s
            guidance_from_c1 = False
            if not guidance_text and c1s and not _looks_like_testing(c1s, req_code):
                guidance_text = c1s
                guidance_from_c1 = True

            # If this row is a "(continued)" line and guidance is in c1,
            # append the leading text to the previous section, then keep any
            # remaining marker-based guidance for normal parsing.
            if guidance_text and guidance_from_c1 and _is_continued_header(c0s):
                section = last_section_by_req.get(req_code)
                if section:
                    mm = MARKER_RE.search(guidance_text)
                    if mm:
                        lead = guidance_text[: mm.start()].strip()
                        if lead:
                            _append_section_continuation(item, section, lead)
                        guidance_text = guidance_text[mm.start() :].strip()
                    else:
                        _append_section_continuation(item, section, guidance_text)
                        guidance_text = ""

            just_set_pending = False
            marker_only = None
            if c0s and c0s.lower() in (m.lower() for m in C0_MARKERS):
                pending_marker_by_table[table_idx] = (c0s, "c0")
                pending_marker_by_req[req_code] = (c0s, "c0")
                just_set_pending = True
            else:
                if c0s and not c1s and not c2s:
                    marker_only = is_marker_only_row(c0s, c0s, "")
                elif c1s and not c0s and not c2s:
                    marker_only = is_marker_only_row(c1s, "", c1s)
                elif c2s and not c0s and not c1s:
                    marker_only = is_marker_only_row(c2s, c2s, "")
            if marker_only and (c0s.lower() == marker_only.lower() or c1s.lower() == marker_only.lower()):
                source_col = "c0" if c0s.lower() == marker_only.lower() else "c1"
                pending_marker_by_table[table_idx] = (marker_only, source_col)
                pending_marker_by_req[req_code] = (marker_only, source_col)
                just_set_pending = True
            elif marker_only and c2s.lower() == marker_only.lower():
                pending_marker_by_table[table_idx] = (marker_only, "c2")
                pending_marker_by_req[req_code] = (marker_only, "c2")
                just_set_pending = True

            pending = pending_marker_by_table.get(table_idx)
            if pending and not just_set_pending:
                pending_marker, source_col = pending
                candidate = c0s if source_col == "c0" else c1s if source_col == "c1" else c2s
                if apply_pending_marker(item, pending_marker, candidate):
                    last_section_by_req[req_code] = {
                        "Customized Approach Objective": "objective",
                        "Applicability Notes": "applicability",
                        "Purpose": "purpose",
                        "Good Practice": "good_practice",
                        "Definitions": "definitions",
                        "Examples": "examples",
                        "Further Information": "further_information",
                    }.get(pending_marker, last_section_by_req.get(req_code, ""))
                    pending_marker_by_table.pop(table_idx, None)
                    pending_marker_by_req.pop(req_code, None)

            if not just_set_pending:
                update_sections(item, c0s, guidance_text, [c0s], req_code, last_section_by_req)
            if not guidance_from_c1:
                update_tests(item, req_code, c1s, pending_marker_by_req)

    for item in data.values():
        dedup_tests(item)

    return data


def write_output(output_path: Path, data: Dict[str, RequirementItem]) -> None:
    lines: List[str] = []
    for code in sorted(data.keys(), key=lambda s: [int(x) for x in s.split(".")]):
        lines.extend(format_requirement(code, data[code]))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def _iter_csv_files(input_dir: Path) -> List[Path]:
    return sorted(p for p in input_dir.iterdir() if p.is_file() and p.suffix.lower() == ".csv")


def _output_path_for(input_csv: Path, output_dir: Path) -> Path:
    return output_dir / f"{input_csv.stem}_formated.txt"


def _docx_output_path_for(input_csv: Path, output_dir: Path) -> Path:
    return output_dir / f"{input_csv.stem}_formated.docx"


def _add_label_paragraphs(doc: Document, label: str, value: str) -> None:
    if not value:
        return
    lines = value.splitlines()
    first = lines[0]
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(first)
    for line in lines[1:]:
        doc.add_paragraph(line)


def _split_requirement_code(code: str) -> Tuple[str, str]:
    parts = code.split(".")
    if len(parts) >= 2:
        control_obj = ".".join(parts[:2])
    else:
        control_obj = code
    return control_obj, code


def _code_lines(code: str) -> List[str]:
    control_obj, sub_req = _split_requirement_code(code)
    return [
        f"Control objectives: {control_obj}",
        f"Sub-requirement: {sub_req}",
    ]


def write_docx(output_path: Path, data: Dict[str, RequirementItem]) -> None:
    doc = Document()
    for idx, code in enumerate(sorted(data.keys(), key=lambda s: [int(x) for x in s.split(".")])):
        item = data[code]
        if idx:
            doc.add_paragraph("---")

        control_obj, sub_req = _split_requirement_code(code)
        p = doc.add_paragraph()
        run = p.add_run("Control objectives: ")
        run.bold = True
        p.add_run(control_obj)
        p = doc.add_paragraph()
        run = p.add_run("Sub-requirement: ")
        run.bold = True
        p.add_run(sub_req)

        if item.requirements:
            _add_label_paragraphs(doc, "Defined Approach Requirements", _bullet_newlines(item.requirements))
        if item.tests:
            if len(item.tests) == 1 and item.tests[0][0] == code:
                _add_label_paragraphs(
                    doc,
                    "Defined Approach Testing Procedures",
                    format_test_desc(item.tests[0][1]),
                )
            else:
                p = doc.add_paragraph()
                run = p.add_run("Defined Approach Testing Procedures:")
                run.bold = True
                for tcode, desc in item.tests:
                    doc.add_paragraph(f'- "{tcode}": {format_test_desc(desc)}')
        if item.objective:
            _add_label_paragraphs(doc, "Customized Approach Objective", _bullet_newlines(item.objective))
        if item.applicability:
            _add_label_paragraphs(doc, "Applicability Notes", _bullet_newlines(item.applicability))
        if item.purpose:
            _add_label_paragraphs(doc, "Guidance - Purpose", _bullet_newlines(item.purpose))
        if item.good_practice:
            _add_label_paragraphs(doc, "Guidance - Good Practice", _bullet_newlines(item.good_practice))
        if item.definitions:
            _add_label_paragraphs(doc, "Guidance - Definitions", _bullet_newlines(item.definitions))
        if item.examples:
            _add_label_paragraphs(doc, "Guidance - Examples", _bullet_newlines(item.examples))
        if item.further_information:
            _add_label_paragraphs(doc, "Guidance - Further Information", _bullet_newlines(item.further_information))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--input-dir",
        default="output/docling",
        help="Directory containing .csv files (default: output/docling)",
    )
    parser.add_argument(
        "--output-dir",
        default="output/format",
        help="Output directory for formatted txt (default: output/format)",
    )
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    if not input_dir.exists():
        raise SystemExit(f"Missing input dir: {input_dir}")
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    csv_files = _iter_csv_files(input_dir)
    if not csv_files:
        raise SystemExit(f"No .csv files found in: {input_dir}")
    for csv_path in csv_files:
        data = parse_csv(csv_path)
        out_path = _output_path_for(csv_path, output_dir)
        write_output(out_path, data)
        print(f"Wrote {out_path}")

        docx_path = _docx_output_path_for(csv_path, output_dir)
        write_docx(docx_path, data)
        print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
