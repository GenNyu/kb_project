#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import sys
from pathlib import Path
from typing import Iterable, Optional

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - runtime dependency check
    print("Thiếu thư viện python-docx. Cài bằng: pip install -r requirements.txt", file=sys.stderr)
    raise SystemExit(2) from exc


def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Chuyển CSV sang Markdown hoặc DOCX.")
    parser.add_argument("input_csv", type=Path, help="File CSV hoặc thư mục chứa CSV.")
    parser.add_argument("output_path", type=Path, help="File .md/.docx hoặc thư mục đích.")
    parser.add_argument(
        "--format",
        choices=["md", "docx"],
        default="md",
        help="Định dạng xuất (md hoặc docx). Mặc định: md.",
    )
    parser.add_argument(
        "--delimiter",
        help="Ký tự phân tách cột. Nếu không cung cấp sẽ tự dò.",
    )
    parser.add_argument(
        "--encoding",
        default="utf-8-sig",
        help="Encoding đọc CSV. Mặc định: utf-8-sig.",
    )
    return parser.parse_args(argv)


def sniff_delimiter(path: Path, encoding: str) -> str:
    sample = path.read_text(encoding=encoding, errors="replace")[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
        return dialect.delimiter
    except csv.Error:
        return ","


def read_csv_rows(path: Path, delimiter: Optional[str], encoding: str) -> list[list[str]]:
    delim = delimiter or sniff_delimiter(path, encoding)
    rows: list[list[str]] = []
    with path.open("r", encoding=encoding, newline="") as handle:
        reader = csv.reader(handle, delimiter=delim)
        for row in reader:
            rows.append([cell.strip() for cell in row])
    return rows


def _normalize_cell(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\n", "<br>")
    text = text.replace("|", "\\|")
    return text


def csv_to_markdown(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    max_cols = max((len(r) for r in rows), default=0)
    if max_cols == 0:
        return []
    normalized = [r + [""] * (max_cols - len(r)) for r in rows]
    header = [_normalize_cell(c) for c in normalized[0]]
    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in normalized[1:]:
        cells = [_normalize_cell(c) for c in row]
        lines.append("| " + " | ".join(cells) + " |")
    return lines


def write_markdown(lines: Iterable[str], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    content = "\n".join(lines).rstrip() + "\n"
    output_path.write_text(content, encoding="utf-8")


def write_docx(rows: list[list[str]], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        Document().save(output_path)
        return
    max_cols = max((len(r) for r in rows), default=0)
    if max_cols == 0:
        Document().save(output_path)
        return
    normalized = [r + [""] * (max_cols - len(r)) for r in rows]
    doc = Document()
    table = doc.add_table(rows=len(normalized), cols=max_cols)
    for r_idx, row in enumerate(normalized):
        for c_idx, cell in enumerate(row):
            table.cell(r_idx, c_idx).text = cell
    doc.save(output_path)


def output_for(input_path: Path, output_path: Path, fmt: str) -> Path:
    if output_path.is_dir() or str(output_path).endswith("/"):
        output_path.mkdir(parents=True, exist_ok=True)
        return output_path / f"{input_path.stem}.{fmt}"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    return output_path


def convert_file(input_path: Path, output_path: Path, fmt: str, delimiter: Optional[str], encoding: str) -> None:
    rows = read_csv_rows(input_path, delimiter, encoding)
    if fmt == "docx":
        write_docx(rows, output_path)
    else:
        lines = csv_to_markdown(rows)
        write_markdown(lines, output_path)


def main(argv: Optional[list[str]] = None) -> int:
    args = parse_args(argv)
    input_path: Path = args.input_csv
    output_path: Path = args.output_path

    if not input_path.exists():
        print(f"Không tìm thấy file hoặc thư mục: {input_path}", file=sys.stderr)
        return 2

    if input_path.is_dir():
        csv_files = sorted(p for p in input_path.iterdir() if p.suffix.lower() == ".csv")
        if not csv_files:
            print("Không tìm thấy file .csv trong thư mục.", file=sys.stderr)
            return 2
        if not output_path.exists():
            output_path.mkdir(parents=True, exist_ok=True)
        for csv_file in csv_files:
            out_file = output_for(csv_file, output_path, args.format)
            convert_file(csv_file, out_file, args.format, args.delimiter, args.encoding)
    else:
        out_file = output_for(input_path, output_path, args.format)
        convert_file(input_path, out_file, args.format, args.delimiter, args.encoding)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
